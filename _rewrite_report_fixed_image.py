# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_MAIN = Path("bao_cao_tfidf_linear_svm_3class_fixed_image.docx")
OUT_FALLBACK = Path("bao_cao_tfidf_linear_svm_3class_fixed_image_updated.docx")
RESULT_DIR = Path("results")


def set_font(run, size=13, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align_center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    set_font(run, size=12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, italic=True)


def add_table(doc, title, headers, rows, widths=None, left_cols=(0,)):
    add_table_title(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row)):
            set_cell_text(cell, value, align_center=idx not in left_cols)
            if widths:
                cell.width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 14, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=12, italic=True, color=(90, 90, 90))


def add_formula(doc, formula):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula)
    set_font(run, size=13, bold=True)


def add_image(doc, image_path, caption, width=6.2):
    image_path = Path(image_path)
    if not image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Chừa chỗ chèn hình]")
        set_font(run, size=12, italic=True, color=(120, 120, 120))
        add_caption(doc, caption)
        return False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)
    return True


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(13)
    for paragraph in doc.paragraphs:
        if paragraph.paragraph_format.space_after is None:
            paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            if run.font.size is None:
                set_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        set_font(run, size=12, bold=run.bold)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO")
    set_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHÂN LOẠI CẢM XÚC ĐÁNH GIÁ NHÀ HÀNG TIẾNG VIỆT BẰNG MÔ HÌNH MACHINE LEARNING")
    set_font(r, size=17, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phương pháp chính: TF-IDF kết hợp mô hình LinearSVC")
    set_font(r, size=13, italic=True)
    doc.add_paragraph()

    add_heading(doc, "1. Giới thiệu", 1)
    add_heading(doc, "1.1. Giới thiệu bài toán", 2)
    add_paragraph(
        doc,
        "Phân loại cảm xúc đánh giá nhà hàng là bài toán tự động xác định thái độ của khách hàng thông qua nội dung đánh giá bằng tiếng Việt. Khi người dùng viết một câu hoặc một đoạn nhận xét về món ăn, phục vụ, giá cả, không gian hoặc trải nghiệm chung, mô hình sẽ dự đoán cảm xúc của đánh giá đó thuộc nhóm không hài lòng, trung lập hoặc hài lòng.",
    )
    add_paragraph(
        doc,
        "Bài toán này có ứng dụng thực tế trong quản lý nhà hàng và phân tích phản hồi khách hàng. Nếu số lượng đánh giá lớn, việc đọc thủ công từng phản hồi sẽ tốn nhiều thời gian. Hệ thống phân loại cảm xúc giúp người quản lý nhanh chóng biết khách hàng đang hài lòng ở điểm nào, chưa hài lòng ở điểm nào và nhóm phản hồi nào cần được ưu tiên xử lý.",
    )
    add_paragraph(
        doc,
        "Trong đề tài này, mô hình nhận đầu vào là một đánh giá nhà hàng bằng tiếng Việt và trả về một trong ba kết quả cảm xúc. Ba lớp được trình bày bằng tiếng Việt để người đọc dễ hiểu, còn trong quá trình xử lý dữ liệu máy tính sẽ dùng các nhãn kỹ thuật tương ứng.",
    )
    add_table(
        doc,
        "Bảng 1. Xác định đầu vào, đầu ra và phạm vi của bài toán",
        ["Thành phần", "Mô tả"],
        [
            ["Số lớp", "3 lớp: Không hài lòng, Trung lập, Hài lòng"],
            ["Đầu vào", "Một câu hoặc một đoạn đánh giá nhà hàng bằng tiếng Việt"],
            ["Đầu ra", "Nhãn cảm xúc dự đoán cho đánh giá đầu vào"],
            ["Mục tiêu", "Phân loại đúng cảm xúc của các đánh giá mới"],
        ],
        widths=[1.8, 5.2],
        left_cols=(1,),
    )
    add_paragraph(
        doc,
        "Bảng 1 cho thấy đề tài được xác định theo hướng phân loại 3 lớp. Việc có lớp trung lập là cần thiết vì nhiều đánh giá nhà hàng thường vừa có điểm khen vừa có điểm chê, ví dụ món ăn ngon nhưng giá cao hoặc phục vụ tốt nhưng không gian chật.",
    )
    add_table(
        doc,
        "Bảng 2. Ý nghĩa của các lớp cảm xúc",
        ["Lớp cảm xúc", "Ý nghĩa"],
        [
            ["Không hài lòng", "Đánh giá có nội dung chê, phàn nàn hoặc thể hiện trải nghiệm chưa tốt."],
            ["Trung lập", "Đánh giá ở mức bình thường hoặc có cả điểm tốt và điểm chưa tốt."],
            ["Hài lòng", "Đánh giá có nội dung khen món ăn, dịch vụ, không gian hoặc trải nghiệm chung."],
        ],
        widths=[1.6, 5.4],
        left_cols=(1,),
    )
    add_paragraph(
        doc,
        "Bảng 2 là cơ sở để đọc các kết quả ở phần sau. Trong ba lớp, lớp trung lập thường khó hơn vì cảm xúc không rõ ràng như hai lớp còn lại.",
    )

    add_heading(doc, "1.2. Giới thiệu công cụ và phương pháp", 2)
    add_paragraph(
        doc,
        "Phần này trình bày các công cụ và phương pháp chính được dùng trong đề tài. Trọng tâm của mô hình gồm hai bước quan trọng: dùng TF-IDF để chuyển đánh giá dạng chữ thành dữ liệu số, sau đó dùng LinearSVC để phân loại dữ liệu số đó vào ba nhóm cảm xúc.",
    )

    add_heading(doc, "1.2.1. Công cụ sử dụng", 2)
    add_paragraph(
        doc,
        "Đề tài được thực hiện bằng Python và thư viện scikit-learn. Python được dùng để đọc dữ liệu, làm sạch văn bản, huấn luyện mô hình và xuất kết quả. scikit-learn cung cấp các công cụ cần thiết cho bài toán học máy truyền thống, gồm biểu diễn văn bản bằng TF-IDF, huấn luyện mô hình phân loại và tính các chỉ số đánh giá.",
    )
    add_table(
        doc,
        "Bảng 3. Công cụ và vai trò trong đề tài",
        ["Công cụ/phương pháp", "Vai trò", "Lý do sử dụng"],
        [
            ["Python", "Xử lý dữ liệu và huấn luyện mô hình", "Dễ dùng, phổ biến trong học máy."],
            ["scikit-learn", "Cung cấp các thuật toán học máy truyền thống", "Phù hợp với bài toán phân loại văn bản."],
            ["TF-IDF", "Chuyển đánh giá dạng chữ thành vector số", "Giúp mô hình nhận biết từ/cụm từ quan trọng."],
            ["LinearSVC", "Phân loại đánh giá vào ba lớp cảm xúc", "Nhanh và hiệu quả với dữ liệu văn bản."],
            ["Ứng dụng web Flask", "Minh họa cách sử dụng mô hình sau khi huấn luyện", "Cho phép nhập một đánh giá hoặc tải file để dự đoán."],
        ],
        widths=[1.8, 2.6, 2.8],
        left_cols=(0, 1, 2),
    )
    add_paragraph(
        doc,
        "Bảng 3 tóm tắt vai trò của từng công cụ. Trọng tâm của đề tài vẫn là mô hình học máy, còn ứng dụng web chỉ đóng vai trò minh họa khả năng triển khai mô hình vào thực tế.",
    )

    add_heading(doc, "1.2.2. Phương pháp TF-IDF", 2)
    add_paragraph(
        doc,
        "TF-IDF có thể hiểu là phương pháp đánh giá mức độ quan trọng của một từ hoặc cụm từ trong một đánh giá. Nếu một từ xuất hiện nhiều trong một đánh giá nhưng không xuất hiện quá phổ biến trong toàn bộ dữ liệu, từ đó thường mang nhiều thông tin hơn. Ví dụ, các cụm như 'rất ngon', 'phục vụ chậm', 'giá cao', 'không quay lại' có thể là dấu hiệu giúp mô hình nhận biết cảm xúc của khách hàng.",
    )
    add_paragraph(doc, "Công thức tính tần suất của một từ trong một đánh giá:")
    add_formula(doc, "TF(t, d) = số lần t xuất hiện trong d / tổng số từ trong d")
    add_paragraph(doc, "Công thức tính mức độ hiếm/phổ biến của từ trong toàn bộ tập dữ liệu:")
    add_formula(doc, "IDF(t) = log((N + 1) / (DF(t) + 1)) + 1")
    add_paragraph(doc, "Công thức kết hợp để tạo giá trị TF-IDF:")
    add_formula(doc, "TF-IDF(t, d) = TF(t, d) × IDF(t)")
    add_paragraph(
        doc,
        "Trong các công thức trên, t là từ hoặc cụm từ đang xét, d là một đánh giá, N là tổng số đánh giá trong tập dữ liệu, còn DF(t) là số đánh giá có chứa t. Giá trị TF-IDF càng cao thì từ hoặc cụm từ đó càng có khả năng mang thông tin quan trọng cho việc phân loại.",
    )
    add_paragraph(
        doc,
        "Trong đề tài, TF-IDF được áp dụng theo hai hướng. Hướng thứ nhất xét theo từ và cụm từ để nắm bắt ý nghĩa trực tiếp của câu đánh giá. Hướng thứ hai xét theo ký tự để mô hình bền hơn với lỗi gõ, viết tắt hoặc biến thể cách viết thường gặp trong đánh giá thực tế. Việc kết hợp hai hướng này giúp mô hình khai thác được nhiều tín hiệu hơn từ văn bản tiếng Việt.",
    )

    add_heading(doc, "1.2.3. Phương pháp LinearSVC", 2)
    add_paragraph(
        doc,
        "LinearSVC là mô hình phân loại dựa trên máy vectơ hỗ trợ tuyến tính. Sau khi TF-IDF biến mỗi đánh giá thành một vector số, LinearSVC học cách tìm ranh giới để tách các nhóm cảm xúc. Với bài toán này, mô hình cần phân biệt ba nhóm: không hài lòng, trung lập và hài lòng.",
    )
    add_paragraph(doc, "Với một đánh giá đã được biểu diễn thành vector x, mô hình tính điểm cho từng lớp cảm xúc theo dạng:")
    add_formula(doc, "Điểm của một lớp = w · x + b")
    add_paragraph(
        doc,
        "Trong đó, x là vector biểu diễn đánh giá sau TF-IDF, w là trọng số mà mô hình học được cho từng đặc trưng và b là hệ số điều chỉnh. Mỗi lớp cảm xúc có một cách tính điểm riêng. Lớp nào có điểm cao nhất sẽ được chọn làm kết quả dự đoán.",
    )
    add_formula(doc, "Kết quả dự đoán = lớp có điểm cao nhất")
    add_paragraph(
        doc,
        "Có thể hiểu đơn giản rằng LinearSVC học xem những từ/cụm từ nào thường gắn với từng cảm xúc. Ví dụ, các tín hiệu như 'ngon', 'nhiệt tình', 'sẽ quay lại' thường hỗ trợ lớp hài lòng; các tín hiệu như 'tệ', 'lâu', 'bẩn', 'không quay lại' thường hỗ trợ lớp không hài lòng. Những đánh giá vừa có khen vừa có chê thường nằm gần ranh giới hơn nên dễ được xếp vào lớp trung lập hoặc bị nhầm với lớp khác.",
    )
    add_paragraph(
        doc,
        "LinearSVC phù hợp với đề tài vì dữ liệu văn bản sau TF-IDF có rất nhiều đặc trưng nhưng mỗi đánh giá chỉ chứa một phần nhỏ trong số đó. Mô hình tuyến tính có tốc độ huấn luyện nhanh, ổn định và thường cho kết quả tốt trong các bài toán phân loại văn bản.",
    )

    add_heading(doc, "1.2.4. Công thức đánh giá mô hình", 2)
    add_paragraph(
        doc,
        "Để đánh giá mô hình, báo cáo không chỉ dùng độ chính xác chung mà còn dùng các chỉ số F1. Lý do là dữ liệu có sự chênh lệch giữa các lớp; nếu chỉ nhìn độ chính xác, mô hình có thể đạt điểm cao nhờ dự đoán tốt lớp có nhiều mẫu nhưng vẫn yếu ở lớp ít mẫu hơn.",
    )
    add_paragraph(doc, "Công thức độ chính xác:")
    add_formula(doc, "Độ chính xác = số dự đoán đúng / tổng số mẫu kiểm tra")
    add_paragraph(doc, "Công thức độ chuẩn xác của một lớp:")
    add_formula(doc, "Độ chuẩn xác = số mẫu dự đoán đúng của lớp đó / tổng số mẫu được dự đoán là lớp đó")
    add_paragraph(doc, "Công thức độ bao phủ của một lớp:")
    add_formula(doc, "Độ bao phủ = số mẫu dự đoán đúng của lớp đó / tổng số mẫu thật thuộc lớp đó")
    add_paragraph(doc, "Công thức F1:")
    add_formula(doc, "F1 = 2 × Độ chuẩn xác × Độ bao phủ / (Độ chuẩn xác + Độ bao phủ)")
    add_paragraph(
        doc,
        "F1 trung bình là trung bình F1 của ba lớp cảm xúc, giúp đánh giá công bằng hơn giữa các lớp. F1 có trọng số cũng dựa trên F1 từng lớp nhưng có xét đến số lượng mẫu của từng lớp, nên phản ánh tốt hơn hiệu quả chung trên toàn bộ tập kiểm tra.",
    )

    add_heading(doc, "2. Thực hiện", 1)
    add_heading(doc, "2.1. Thu thập dữ liệu", 2)
    add_paragraph(
        doc,
        "Dữ liệu sử dụng trong đề tài là bộ dữ liệu đánh giá nhà hàng tiếng Việt. Mỗi mẫu dữ liệu gồm nội dung đánh giá và nhãn cảm xúc tương ứng. Bộ dữ liệu đã được chia thành ba phần: tập huấn luyện dùng để dạy mô hình, tập thẩm định dùng để chọn phương án phù hợp và tập kiểm tra cuối dùng để đánh giá kết quả sau cùng.",
    )
    add_paragraph(
        doc,
        "Dữ liệu được thu thập từ các phản hồi của người dùng về nhà hàng, quán ăn hoặc dịch vụ ăn uống. Nội dung đánh giá có thể liên quan đến chất lượng món ăn, thái độ phục vụ, giá cả, không gian, thời gian chờ hoặc trải nghiệm tổng thể. Trước khi huấn luyện, dữ liệu được kiểm tra và làm sạch để giảm trùng lặp, giảm nhiễu và hạn chế việc kết quả đánh giá bị sai lệch.",
    )
    add_table(
        doc,
        "Bảng 4. Phân bố dữ liệu gốc theo từng tập",
        ["Tập dữ liệu", "Tổng số mẫu", "Không hài lòng", "Trung lập", "Hài lòng"],
        [
            ["Huấn luyện", "26,170", "4,974", "5,347", "15,849"],
            ["Thẩm định", "3,271", "622", "668", "1,981"],
            ["Kiểm tra", "3,272", "622", "669", "1,981"],
            ["Tổng", "32,713", "6,218", "6,684", "19,811"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 4 cho thấy số lượng đánh giá hài lòng lớn hơn hai nhóm còn lại. Đây là đặc điểm thường gặp trong dữ liệu đánh giá nhà hàng. Nếu không chú ý đến sự chênh lệch này, mô hình có thể có xu hướng dự đoán nhiều hơn về lớp có số lượng mẫu lớn.",
    )
    add_table(
        doc,
        "Bảng 5. Số lượng dữ liệu sau khi làm sạch",
        ["Tập dữ liệu", "Trước làm sạch", "Sau làm sạch", "Không hài lòng", "Trung lập", "Hài lòng"],
        [
            ["Huấn luyện", "26,170", "25,640", "4,902", "5,244", "15,494"],
            ["Thẩm định", "3,271", "3,165", "605", "643", "1,917"],
            ["Kiểm tra", "3,272", "3,137", "602", "638", "1,897"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 5 trình bày số lượng dữ liệu sau khi loại bỏ các mẫu gây nhiễu. Sau bước này, tập kiểm tra cuối còn 3,137 đánh giá và được giữ riêng để đánh giá mô hình sau khi đã hoàn tất quá trình lựa chọn phương án.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_class_distribution.png",
        "Hình 1. Phân bố dữ liệu theo ba lớp cảm xúc.",
        width=6.2,
    )
    add_paragraph(
        doc,
        "Hình 1 minh họa trực quan sự chênh lệch giữa các lớp dữ liệu. Lớp hài lòng chiếm tỷ lệ cao nhất, trong khi hai lớp không hài lòng và trung lập có số lượng thấp hơn. Vì vậy, khi huấn luyện cần quan tâm đến khả năng nhận diện cả ba lớp, không chỉ nhìn vào độ chính xác chung.",
    )
    add_table(
        doc,
        "Bảng 6. Một số ví dụ đánh giá trong dữ liệu",
        ["Ví dụ đánh giá", "Lớp cảm xúc"],
        [
            ["Phục vụ lâu, phải chờ để có menu, cốc trà bẩn, thái độ nhân viên làm cho có.", "Không hài lòng"],
            ["Món ăn ngon, chất lượng ổn định, tuy nhiên phục vụ chưa chuyên nghiệp và bàn ghế chưa hợp lý.", "Trung lập"],
            ["Đồ ăn ngon, nhân viên phục vụ rất nhiệt tình và chu đáo.", "Hài lòng"],
        ],
        widths=[5.4, 1.6],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 6 cho thấy đặc điểm của từng lớp. Đánh giá không hài lòng thường chứa các từ thể hiện phàn nàn. Đánh giá hài lòng có tín hiệu khen rõ ràng. Đánh giá trung lập thường có cả nhận xét tốt và chưa tốt nên khó phân loại hơn.",
    )

    add_heading(doc, "2.2. Huấn luyện mô hình", 2)
    add_paragraph(
        doc,
        "Quá trình huấn luyện bắt đầu bằng việc chuẩn hóa nội dung đánh giá. Văn bản được chuyển về dạng thống nhất, loại bỏ các thành phần không cần thiết như đường dẫn trang web hoặc ký tự thừa, đồng thời giữ lại dấu tiếng Việt vì dấu có ảnh hưởng trực tiếp đến nghĩa của từ. Sau đó, nội dung đánh giá được chuyển thành vector số bằng TF-IDF.",
    )
    add_paragraph(
        doc,
        "Sau khi có đặc trưng TF-IDF, các mô hình phân loại được huấn luyện và so sánh. Đề tài thử một số mô hình học máy truyền thống, trong đó mô hình SVM tuyến tính thông qua LinearSVC cho kết quả tổng thể tốt nhất. Để người đọc dễ theo dõi, báo cáo trình bày các phương án bằng tên mô tả thay vì dùng tên cấu hình trong code.",
    )
    add_table(
        doc,
        "Bảng 7. Các bước chính trong quá trình huấn luyện",
        ["Bước", "Cách thực hiện", "Ý nghĩa"],
        [
            ["Chuẩn hóa văn bản", "Làm sạch và thống nhất nội dung đánh giá", "Giảm nhiễu trước khi đưa vào mô hình."],
            ["Biểu diễn văn bản", "Dùng TF-IDF theo từ/cụm từ và theo ký tự", "Chuyển văn bản thành số để máy tính xử lý."],
            ["Huấn luyện mô hình", "Dùng các mô hình học máy truyền thống", "Cho mô hình học mối liên hệ giữa đánh giá và cảm xúc."],
            ["So sánh phương án", "Đánh giá trên tập thẩm định", "Chọn phương án có khả năng tổng quát tốt hơn."],
            ["Đánh giá cuối", "Kiểm tra trên tập dữ liệu riêng", "Đo hiệu quả mô hình trên dữ liệu chưa dùng để lựa chọn."],
        ],
        widths=[1.6, 2.7, 2.7],
        left_cols=(0, 1, 2),
    )
    add_paragraph(
        doc,
        "Bảng 7 mô tả toàn bộ quy trình huấn luyện ở mức dễ hiểu. Điểm quan trọng là tập kiểm tra cuối không được dùng để chọn mô hình, nhằm bảo đảm kết quả đánh giá khách quan hơn.",
    )
    add_table(
        doc,
        "Bảng 8. So sánh một số phương án huấn luyện trên tập thẩm định",
        ["Phương án", "Độ chính xác", "F1 trung bình", "F1 có trọng số", "F1 lớp không hài lòng", "F1 lớp trung lập", "F1 lớp hài lòng"],
        [
            ["SVM tuyến tính, tăng hỗ trợ nhẹ cho lớp trung lập", "0.8575", "0.8211", "0.8564", "0.8863", "0.6661", "0.9108"],
            ["Hồi quy Logistic, tăng hỗ trợ nhẹ cho lớp trung lập", "0.8559", "0.8205", "0.8558", "0.8845", "0.6672", "0.9100"],
            ["SVM tuyến tính, bổ sung thêm mẫu trung lập", "0.8562", "0.8204", "0.8557", "0.8850", "0.6661", "0.9100"],
            ["Hồi quy Logistic, bổ sung thêm mẫu trung lập", "0.8543", "0.8208", "0.8550", "0.8865", "0.6682", "0.9078"],
            ["SVM tuyến tính, hỗ trợ trung lập ở mức thấp hơn", "0.8566", "0.8190", "0.8551", "0.8848", "0.6619", "0.9105"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 8 cho thấy các phương án có kết quả khá gần nhau. Phương án SVM tuyến tính có hỗ trợ nhẹ cho lớp trung lập đạt kết quả tổng thể tốt nhất trên tập thẩm định, nên được chọn làm mô hình cuối. Cách trình bày này tập trung vào ý nghĩa của phương án thay vì tên biến hoặc tham số trong chương trình.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_validation_model_comparison.png",
        "Hình 2. So sánh kết quả các phương án huấn luyện trên tập thẩm định.",
        width=6.2,
    )
    add_paragraph(
        doc,
        "Hình 2 minh họa trực quan kết quả ở Bảng 8. Mục đích của hình là cho thấy mô hình được chọn không dựa trên cảm tính mà dựa trên so sánh định lượng giữa nhiều phương án khác nhau.",
    )

    add_heading(doc, "2.3. Kiểm tra mô hình", 2)
    add_paragraph(
        doc,
        "Sau khi chọn được mô hình cuối, mô hình được kiểm tra trên tập dữ liệu riêng gồm 3,137 đánh giá. Tập này không dùng trong quá trình lựa chọn mô hình. Do đó, kết quả ở phần này được xem là kết quả chính để nhận xét chất lượng mô hình.",
    )
    add_paragraph(
        doc,
        "Các chỉ số được sử dụng gồm độ chính xác, F1 trung bình và F1 có trọng số. Độ chính xác cho biết tỷ lệ dự đoán đúng trên toàn bộ tập kiểm tra. F1 trung bình giúp xem xét công bằng hơn giữa các lớp. F1 có trọng số phản ánh hiệu quả chung có xét đến số lượng mẫu của từng lớp.",
    )
    add_table(
        doc,
        "Bảng 9. Kết quả tổng hợp trên ba tập dữ liệu",
        ["Tập dữ liệu", "Độ chính xác", "F1 trung bình", "F1 có trọng số", "F1 không hài lòng", "F1 trung lập", "F1 hài lòng"],
        [
            ["Huấn luyện", "0.9940", "0.9932", "0.9941", "0.9980", "0.9860", "0.9955"],
            ["Thẩm định", "0.8575", "0.8211", "0.8564", "0.8863", "0.6661", "0.9108"],
            ["Kiểm tra cuối", "0.8406", "0.8038", "0.8417", "0.8614", "0.6503", "0.8998"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 9 cho thấy mô hình đạt độ chính xác 0.8406 trên tập kiểm tra cuối. Kết quả trên tập huấn luyện cao hơn rõ rệt so với hai tập còn lại, cho thấy mô hình có học rất tốt dữ liệu đã thấy. Tuy nhiên, kết quả trên tập thẩm định và tập kiểm tra cuối vẫn tương đối gần nhau, nên mô hình có khả năng áp dụng cho dữ liệu mới ở mức chấp nhận được.",
    )
    add_table(
        doc,
        "Bảng 10. Kết quả chi tiết theo từng lớp trên tập kiểm tra cuối",
        ["Lớp cảm xúc", "Độ chuẩn xác", "Độ bao phủ", "F1", "Số mẫu"],
        [
            ["Không hài lòng", "0.8658", "0.8571", "0.8614", "602"],
            ["Trung lập", "0.6366", "0.6646", "0.6503", "638"],
            ["Hài lòng", "0.9051", "0.8946", "0.8998", "1,897"],
            ["Trung bình", "0.8025", "0.8054", "0.8038", "3,137"],
            ["Trung bình có xét số mẫu", "0.8429", "0.8406", "0.8417", "3,137"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 10 cho thấy mô hình nhận diện tốt nhất ở lớp hài lòng, tiếp theo là lớp không hài lòng. Lớp trung lập có kết quả thấp hơn vì nội dung đánh giá thường không rõ ràng, vừa có yếu tố tích cực vừa có yếu tố tiêu cực.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_final_test_class_metrics.png",
        "Hình 3. Kết quả theo từng lớp cảm xúc trên tập kiểm tra cuối.",
        width=6.2,
    )
    add_paragraph(
        doc,
        "Hình 3 giúp quan sát nhanh sự khác biệt giữa ba lớp. Cột của lớp trung lập thấp hơn hai lớp còn lại, đây là hạn chế chính của mô hình hiện tại.",
    )
    add_table(
        doc,
        "Bảng 11. Ma trận nhầm lẫn của mô hình trên tập kiểm tra cuối",
        ["Nhãn thật / Dự đoán", "Không hài lòng", "Trung lập", "Hài lòng"],
        [
            ["Không hài lòng", "516", "74", "12"],
            ["Trung lập", "48", "424", "166"],
            ["Hài lòng", "32", "168", "1,697"],
        ],
        left_cols=(0,),
    )
    add_paragraph(
        doc,
        "Bảng 11 cho biết mô hình dự đoán đúng nhiều nhất ở các ô nằm trên đường chéo chính: 516 đánh giá không hài lòng, 424 đánh giá trung lập và 1,697 đánh giá hài lòng. Các ô còn lại là những trường hợp mô hình dự đoán nhầm.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_final_test_confusion_matrix_counts.png",
        "Hình 4. Ma trận nhầm lẫn dạng số lượng.",
        width=5.6,
    )
    add_paragraph(
        doc,
        "Hình 4 cho thấy phần lớn đánh giá hài lòng được nhận diện đúng. Lỗi đáng chú ý nhất nằm giữa lớp trung lập và lớp hài lòng, vì nhiều đánh giá trung lập vẫn chứa các từ khen như 'ngon', 'ổn', 'được' nhưng đi kèm một điểm chưa tốt.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_final_test_confusion_matrix_normalized.png",
        "Hình 5. Ma trận nhầm lẫn theo tỷ lệ.",
        width=5.6,
    )
    add_paragraph(
        doc,
        "Hình 5 biểu diễn kết quả theo tỷ lệ để so sánh công bằng hơn giữa các lớp. Kết quả này tiếp tục cho thấy lớp trung lập là lớp khó nhất và cần được cải thiện nếu phát triển đề tài thêm.",
    )
    add_image(
        doc,
        RESULT_DIR / "final_overall_final_test_error_pairs.png",
        "Hình 6. Các nhóm nhầm lẫn xuất hiện nhiều nhất.",
        width=6.2,
    )
    add_paragraph(
        doc,
        "Hình 6 cho thấy hai nhóm nhầm lẫn lớn nhất là đánh giá hài lòng bị dự đoán thành trung lập và đánh giá trung lập bị dự đoán thành hài lòng. Đây là kết quả phù hợp với đặc điểm dữ liệu, vì ranh giới giữa hai nhóm này thường không rõ như ranh giới giữa không hài lòng và hài lòng.",
    )

    add_heading(doc, "3. Kết luận", 1)
    add_paragraph(
        doc,
        "Đề tài đã xây dựng được mô hình phân loại cảm xúc đánh giá nhà hàng tiếng Việt theo hướng học máy truyền thống. Mô hình sử dụng TF-IDF để biểu diễn văn bản và LinearSVC để phân loại cảm xúc. Kết quả trên tập kiểm tra cuối đạt độ chính xác 0.8406, F1 trung bình 0.8038 và F1 có trọng số 0.8417.",
    )
    add_paragraph(
        doc,
        "Về độ chính xác, mô hình hoạt động tốt với hai lớp hài lòng và không hài lòng. Lớp trung lập còn khó do nội dung đánh giá thường pha trộn cả ý khen và ý chê. Về sự thuận tiện, phương pháp TF-IDF kết hợp LinearSVC có thời gian huấn luyện nhanh, dễ triển khai và phù hợp với phạm vi đồ án. Mô hình cũng đã được tích hợp vào ứng dụng web minh họa để nhập đánh giá trực tiếp hoặc phân tích file dữ liệu.",
    )
    add_table(
        doc,
        "Bảng 12. Tổng hợp nhận xét về mô hình",
        ["Nội dung", "Nhận xét"],
        [
            ["Độ chính xác", "Đạt 0.8406 trên tập kiểm tra cuối; lớp hài lòng và không hài lòng có kết quả tốt hơn lớp trung lập."],
            ["Sự thuận tiện", "Quy trình huấn luyện và triển khai tương đối đơn giản, có thể dùng trong ứng dụng web minh họa."],
            ["Thời gian", "Huấn luyện nhanh, không yêu cầu tài nguyên tính toán lớn."],
            ["Ưu điểm", "Dễ hiểu, dễ giải thích, phù hợp với bài toán phân loại văn bản tiếng Việt."],
            ["Khuyết điểm", "Chưa xử lý thật tốt các đánh giá trung lập hoặc đánh giá có nhiều ý trái chiều."],
            ["Hướng phát triển", "Thu thập thêm dữ liệu, bổ sung nhiều đánh giá trung lập hơn, cải thiện tiền xử lý tiếng Việt và triển khai thành ứng dụng web/mobile hoàn chỉnh."],
        ],
        widths=[1.7, 5.4],
        left_cols=(0, 1),
    )
    add_paragraph(
        doc,
        "Ưu điểm chính của phương pháp là đơn giản, dễ giải thích và phù hợp với yêu cầu sử dụng mô hình học máy truyền thống. Tuy nhiên, mô hình vẫn còn hạn chế khi gặp các đánh giá có cảm xúc không rõ ràng, đặc biệt là nhóm trung lập. Trong tương lai, có thể cải thiện bằng cách thu thập thêm dữ liệu thực tế, chuẩn hóa cách gán nhãn và mở rộng ứng dụng để hỗ trợ người dùng phân tích đánh giá thuận tiện hơn.",
    )

    style_doc(doc)
    return doc


doc = build_report()
try:
    doc.save(OUT_MAIN)
    print(OUT_MAIN.resolve())
except PermissionError:
    doc.save(OUT_FALLBACK)
    print(OUT_FALLBACK.resolve())
